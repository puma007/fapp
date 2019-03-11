# -*- coding: utf-8 -*-
"""
Created on 2019-03-09
成绩分析程序
@author: hujili
"""

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
from scipy.stats import kstest
from docx import Document
from docx.shared import Inches


def normfun(x, mu, sigma):
    """正态分布函数
        Args:
            mu: 均值.
            sigma:标准差.
            pdf:概率密度函数.
            np.exp():概率密度函数公式
        Returns:
            pdf:概率密度函数
        """
    pdf = np.exp(-((x - mu) ** 2) / (2 * sigma ** 2)) / (sigma * np.sqrt(2 * np.pi))
    return pdf


def analysis(xlsfile, docname, filename):
    """成绩分析函数
            Args:
                xlsfile： 电子表格文件
                docname： 模板文件名称 成绩分析表模板.docx
                filename：文件名称

    """
    res = {"msg": None, "code": 0}
    # 读取电子表格
    try:
        xls_data = pd.read_excel(xlsfile, sheet_name=0)
        # 获取sheet1中期末成绩列数据
        score_data = xls_data['期末成绩'].dropna()
    except Exception:
        res.update(msg="读取数据错误,未找到期末成绩列",code=-1)
        return res
    else:
        print(len(score_data))
        print(score_data)
    # 计算均值和方差
    try:
        mean = score_data.mean()
        std = score_data.std()
    except Exception:
        res.update(msg="数据类型错误，请删除旷缺等非数值行", code=-1)
        return res
    else:
        print('均值：', mean, "方差：", std)
    # x的范围为30-100，以1为单位
    x = np.arange(20, 100, 1)
    # x数对应的概率密度
    y = normfun(x, mean, std)
    print("概率密度：")
    print(y)
    # 绘图 参数,颜色，线宽
    plt.clf()
    plt.plot(x, y, color='g', linewidth=3)
    # 柱状图 数据，数组，颜色，颜色深浅，组宽，显示频率
    plt.hist(score_data.dropna(), bins=7, color='r', alpha=0.5, rwidth=0.9, normed=True)
    plt.title('exam analysis')
    plt.xlabel('Final exam', fontproperties='Courier')
    plt.ylabel('Frequency')
    plt.savefig(os.path.splitext(filename)[0] + ".png")
    # plt.show()
    # ks检验判断
    cols = score_data.copy(deep=True)
    for i in range(len(score_data)):
        print(i)
        cols[i] = (score_data[i] - mean) / std
    (d, p) = kstest(cols, "norm", N=len(score_data), mode="asymp")
    print("ks=")
    print(p)
    # 结果写入doc文件,word版本要>2007
    return save_docx(docname, score_data, filename, p)


def save_docx(docname, score_data, filename, pvalue):
    """分析结果写入doc模板函数
                Args:
                    docname： 模板文件名称 成绩分析表模板.docx
                    score_data：数据
                    filename：文件名称
                    pvalue:ks检验，Asymp. Sig. (2-tailed)值
    """
    res = {"msg": None, "code": 0}
    # 读取成绩分析表word文件
    document = Document(docname)
    # 写入图片
    p = document.tables[0].rows[17].cells[1].add_paragraph()
    r = p.add_run()
    r.add_picture(os.path.splitext(filename)[0] + ".png", width=Inches(3.36), height=Inches(2.52))

    # 统计各个分数段人数
    # >=90分的人数
    n90 = len((np.where(score_data >= 90))[0])
    # >=80分的人数
    n80 = len((np.where(score_data >= 80))[0]) - n90
    # >=70分的人数
    n70 = len((np.where(score_data >= 70))[0]) - n80 - n90
    # >=60分的人数
    n60 = len((np.where(score_data >= 60))[0]) - n80 - n90 - n70
    # <60分的人数
    n0 = len((np.where(score_data >= 0))[0]) - n90 - n80 - n70 - n60
    # 总人数
    total = n90 + n80 + n70 + n60 + n0
    print('各个分数段人数,总数：', n90, n80, n70, n60, n0, ',', total)

    # 填充考试人数
    document.tables[0].rows[17].cells[6].text = str(total)
    # 填充平均分
    document.tables[0].rows[17].cells[8].text = str(round(score_data.mean(), 1))
    # 填充最高分
    document.tables[0].rows[18].cells[6].text = str(score_data.max())
    # 填充最低分
    document.tables[0].rows[18].cells[8].text = str(score_data.min())

    # ≥90分占比
    print("≥90分占比 %.2f%%" % (n90 / total * 100))
    document.tables[0].rows[19].cells[6].text = "%d 人,占%.2f%%" % (n90, n90 / total * 100)
    # 80～89分占比
    print("0～89分占比 %.2f%%" % (n80 / total * 100))
    document.tables[0].rows[20].cells[6].text = "%d 人,占%.2f%%" % (n80, n80 / total * 100)
    # 70～79分占比
    print("70～79分占比 %.2f%%" % (n70 / total * 100))
    document.tables[0].rows[21].cells[6].text = "%d 人,占%.2f%%" % (n70, n70 / total * 100)
    # 60～69分占比
    print("60～69分占比 %.2f%%" % (n60 / total * 100))
    document.tables[0].rows[22].cells[6].text = "%d 人,占%.2f%%" % (n60, n60 / total * 100)
    # ＜60分占比
    print("＜60分占比 %.2f%%" % (n0 / total * 100))
    document.tables[0].rows[23].cells[6].text = "%d 人,占%.2f%%" % (n0, n0 / total * 100)
    # 写入试卷难易度分析
    if pvalue > 0.05:
        document.tables[1].rows[0].cells[
            1].text = "试卷难易度分析：\n 1.经非参数1-sample-K-S检验，Asymp. Sig. (2-tailed) %.5f>0.05；学生考试成绩符合正态分布。\n 2.学生的考试成绩平均分=%.1f ，本试卷难易程度适中，题量适中." % (
            pvalue,score_data.mean())
    else:
        document.tables[1].rows[0].cells[
            1].text = "试卷难易度分析：\n 1.经非参数1-sample-K-S检验，Asymp. Sig. (2-tailed) %.5f<=0.05；学生考试成绩不符合正态分布。\n 2.学生的考试成绩平均分=%.1f ，本试卷难易程度适中，题量适中." % (
            pvalue,score_data.mean())
    # 分析结果保存到同名的docx文件中
    try:
        document.save(os.path.splitext(filename)[0] + '.docx')
    except Exception:
        res.update(msg="生成word文档失败",code=-1)
        return res
    else:
        print('分析结果已自动写入：' + os.path.splitext(filename)[0] + '.docx')
    return res


if __name__ == '__main__':
    # 获取当前目录
    path = os.getcwd()
    dirs = os.listdir(path)
    # 遍历当前目录下的xlsx文件
    for f in dirs:
        if os.path.splitext(f)[1] == ".xlsx" or os.path.splitext(f)[1] == ".xls":
            print('正在分析：' + f + '文件')
            # 成绩分析
            analysis(f,"成绩分析模板.docx",f)
